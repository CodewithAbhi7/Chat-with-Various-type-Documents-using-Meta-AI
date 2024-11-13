import streamlit as st
import docx
from bs4 import BeautifulSoup
from PIL import Image
import pytesseract
from PyPDF2 import PdfReader
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.text_splitter import CharacterTextSplitter
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from meta_ai_api import MetaAI  # Importing Meta AI API

# Initialize Meta AI
ai = MetaAI()

def get_pdf_text(pdf_docs):
    """Extract text from PDF files."""
    text = ""
    for pdf in pdf_docs:
        pdf_reader = PdfReader(pdf)
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def get_docx_text(docx_files):
    """Extract text from DOCX files."""
    text = ""
    for docx_file in docx_files:
        doc = docx.Document(docx_file)
        # Extract paragraphs and tables
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
    return text

def get_html_text(html_files):
    """Extract text from HTML files."""
    text = ""
    for html_file in html_files:
        content = html_file.read()  # Ensure file content is read properly
        soup = BeautifulSoup(content, 'html.parser')
        extracted_text = soup.get_text(separator="\n").strip()
        if extracted_text:
            text += extracted_text + "\n"
    return text

def get_image_text(image_files):
    """Extract text from images using Tesseract OCR."""
    text = ""
    for image_file in image_files:
        try:
            img = Image.open(image_file)
            extracted_text = pytesseract.image_to_string(img)
            if extracted_text.strip():
                text += extracted_text + "\n"
        except Exception as e:
            st.error(f"Error extracting text from image: {e}")
    return text

def get_text_chunks(text):
    """Split text into smaller, meaningful chunks."""
    text_splitter = RecursiveCharacterTextSplitter(
        separators=["\n\n", "\n", ".", " "],  # Try splitting by paragraphs, lines, or sentences
        chunk_size=512,  # Adjust as per need
        chunk_overlap=50,  # Overlap to maintain context across chunks
    )
    chunks = text_splitter.split_text(text)

    # Ensure non-empty and meaningful chunks
    valid_chunks = [chunk.strip() for chunk in chunks if chunk.strip()]
    if not valid_chunks:
        st.warning("No valid chunks created. Try adjusting the chunk size or input text.")
    
    return valid_chunks

def get_vectorstore(text_chunks):
    """Create a FAISS vector store from text chunks.""" 
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    vectorstore = FAISS.from_texts(text_chunks, embeddings)
    return vectorstore

def process_files(file_types, pdf_docs, docx_files, html_files, image_files):
    """Process multiple files and create vectorstore."""
    text = ""

    if "PDF" in file_types and pdf_docs:
        text += get_pdf_text(pdf_docs)
    if "DOCX" in file_types and docx_files:
        text += get_docx_text(docx_files)
    if "HTML" in file_types and html_files:
        text += get_html_text(html_files)
    if "Image" in file_types and image_files:
        text += get_image_text(image_files)

    if text:
        st.session_state.document_text = text  # Store document text
        st.write("Successfully extracted text from files.")
        
        # Create text chunks
        text_chunks = get_text_chunks(text)
        if text_chunks:
            st.write(f"Created {len(text_chunks)} text chunks.")
            return get_vectorstore(text_chunks)
        else:
            st.warning("Failed to create valid text chunks.")
            return None
    else:
        st.warning("No text extracted from the uploaded files.")
        return None

import logging

# Set up logging
logging.basicConfig(level=logging.DEBUG)

def handle_userinput(user_question):
    """Handle user input for Q&A using Meta AI API."""
    if "document_text" not in st.session_state or not st.session_state.document_text:
        st.warning("No document text available. Please upload files and process them first.")
        return

    context = st.session_state.document_text  
    full_input = f"Context: {context}\n\nUser question: {user_question}"

    logging.debug(f"Sending request to Meta AI with input: {full_input}")

    try:
        response = ai.prompt(message=full_input)
    except Exception as e:
        logging.error(f"Error occurred: {e}")
        st.error("An error occurred while trying to get a response from Meta AI.")
        return

    # Extract the message text from the response
    answer = response.get('message', 'No answer provided.')

    # Clean and display the answer without code-like formatting
    st.session_state.chat_history.append((user_question, answer))

    # Display chat history in plain text
    for q, a in st.session_state.chat_history:
        st.write(f"**Q:** {q}")
        st.text_area("A:", a, height=200)

def main():
    """Main Streamlit app function.""" 
    st.header("Chat with Documents using Meta AI")

    # Initialize session state attributes
    if 'conversation' not in st.session_state:
        st.session_state.conversation = None
        st.session_state.chat_history = []
        st.session_state.document_text = ""  # Initialize document_text to an empty string

    user_question = st.text_input("Ask a question:")

    if user_question:
        handle_userinput(user_question)

    # File uploader and types selection
    file_types = st.multiselect("Select file types:", ["PDF", "DOCX", "HTML", "Image"])
    
    # File uploaders for different document types
    pdf_docs = st.file_uploader("Upload PDFs", accept_multiple_files=True, type="pdf")
    docx_files = st.file_uploader("Upload DOCX files", accept_multiple_files=True, type="docx")
    html_files = st.file_uploader("Upload HTML files", accept_multiple_files=True, type="html")
    image_files = st.file_uploader("Upload Images", accept_multiple_files=True, type=["jpg", "jpeg", "png"])

    # Process button
    if st.button("Process"):
        with st.spinner("Processing..."):
            vectorstore = process_files(file_types, pdf_docs, docx_files, html_files, image_files)
            if vectorstore:
                st.success("Documents processed successfully! Ask your questions.")
            else:
                st.error("Failed to process the uploaded documents.")

if __name__ == '__main__':
    main()